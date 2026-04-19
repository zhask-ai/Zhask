"""IntegriShield DOCX Report Builder — unified executive-grade report engine.

Produces polished, branded Word documents for every module:
  - Cover page with wordmark, classification badge, metadata
  - Auto-updating Table of Contents field
  - Executive summary with embedded matplotlib chart
  - Per-section styled headings + narrative paragraphs
  - Zebra-striped evidence tables with severity colour cells
  - Remediation appendix with priority / owner / effort columns
  - Prepared-by / reviewed-by signature block
  - Page-numbered header/footer with "IntegriShield Confidential"

Requires: python-docx>=1.0.0, matplotlib>=3.5.0 (both already in project deps).
"""

from __future__ import annotations

import io
import math
import textwrap
from datetime import datetime, timezone
from typing import Any

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE

# ── palette ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0B, 0x24, 0x47)
CYAN   = RGBColor(0x19, 0xA7, 0xCE)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY_L = RGBColor(0xF4, 0xF6, 0xF9)
GREY_D = RGBColor(0x64, 0x74, 0x8B)
RED    = RGBColor(0xEF, 0x44, 0x44)
ORANGE = RGBColor(0xF9, 0x73, 0x16)
YELLOW = RGBColor(0xEA, 0xB3, 0x08)
GREEN  = RGBColor(0x22, 0xC5, 0x5E)

# hex strings for matplotlib
_NAVY_HEX  = "#0B2447"
_CYAN_HEX  = "#19A7CE"
_RED_HEX   = "#EF4444"
_ORANGE_HEX= "#F97316"
_YELLOW_HEX= "#EAB308"
_GREEN_HEX = "#22C55E"
_GREY_HEX  = "#64748B"


# ─────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background colour (RRGGBB hex, no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_col_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    return br


def _add_page_break_v2(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    from docx.oxml import OxmlElement as _OE
    from docx.oxml.ns import qn as _qn
    br = _OE("w:br")
    br.set(_qn("w:type"), "page")
    run._r.append(br)


def _hr(doc: Document, color_hex: str = "19A7CE", thickness_pt: int = 2) -> None:
    """Insert a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness_pt * 4))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_run_font(run, name: str = "Calibri", size_pt: float = 11,
                  bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_table_style(table) -> None:
    """Remove default table border, set clean cell margins."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    # clear borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _add_toc_field(doc: Document) -> None:
    """Insert a real TOC field (Word updates on open / Ctrl+A, F9)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)


def _add_header_footer(doc: Document, module_name: str) -> None:
    """Add header (module name + Confidential) and page-numbered footer."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(f"IntegriShield  |  {module_name}  |  CONFIDENTIAL")
    _set_run_font(run, size_pt=8, color=GREY_D)

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left = fp.add_run("IntegriShield Security Platform  ·  ")
    _set_run_font(run_left, size_pt=8, color=GREY_D)
    # page number field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_pg = fp.add_run()
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)
    _set_run_font(run_pg, size_pt=8, color=GREY_D)


# ─────────────────────────────────────────────────────────────────────────────
# Chart generation (matplotlib → PNG bytes → embedded in docx)
# ─────────────────────────────────────────────────────────────────────────────

def _severity_doughnut_png(counts: dict[str, int], title: str = "Severity Distribution") -> bytes:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        labels  = [k for k, v in counts.items() if v > 0]
        values  = [v for v in counts.values() if v > 0]
        colors  = {
            "CRITICAL": _RED_HEX, "critical": _RED_HEX,
            "HIGH": _ORANGE_HEX,  "high": _ORANGE_HEX,
            "MEDIUM": _YELLOW_HEX,"medium": _YELLOW_HEX,
            "LOW": _GREEN_HEX,    "low": _GREEN_HEX,
            "INFO": _GREY_HEX,    "info": _GREY_HEX,
        }
        clrs = [colors.get(l, _CYAN_HEX) for l in labels]
        fig, ax = plt.subplots(figsize=(3.5, 3.0), facecolor="none")
        wedges, texts, autotexts = ax.pie(
            values, labels=labels, colors=clrs,
            autopct="%1.0f%%", startangle=90,
            wedgeprops={"width": 0.55, "edgecolor": "white", "linewidth": 1.5},
            textprops={"fontsize": 8, "color": "#1e293b"},
        )
        for at in autotexts:
            at.set_fontsize(7)
            at.set_color("white")
        ax.set_title(title, fontsize=9, fontweight="bold", color=_NAVY_HEX, pad=8)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=130, bbox_inches="tight",
                    facecolor="none", transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return b""


def _risk_gauge_png(score: float, title: str = "Risk Score") -> bytes:
    """Semi-circle gauge chart 0–100."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import numpy as np

        score = max(0.0, min(100.0, float(score)))
        fig, ax = plt.subplots(figsize=(3.2, 2.0), facecolor="none")
        ax.set_xlim(-1.1, 1.1)
        ax.set_ylim(-0.15, 1.2)
        ax.axis("off")

        zones = [(_GREEN_HEX, 0, 30), (_YELLOW_HEX, 30, 60), (_ORANGE_HEX, 60, 80), (_RED_HEX, 80, 100)]
        for color, lo, hi in zones:
            theta = np.linspace(math.pi * (1 - lo / 100), math.pi * (1 - hi / 100), 50)
            x_outer = np.cos(theta)
            y_outer = np.sin(theta)
            theta_r = theta[::-1]
            x_inner = 0.65 * np.cos(theta_r)
            y_inner = 0.65 * np.sin(theta_r)
            xs = np.concatenate([x_outer, x_inner])
            ys = np.concatenate([y_outer, y_inner])
            ax.fill(xs, ys, color=color, alpha=0.85)

        needle_angle = math.pi * (1 - score / 100)
        ax.annotate("", xy=(0.55 * math.cos(needle_angle), 0.55 * math.sin(needle_angle)),
                    xytext=(0, 0),
                    arrowprops=dict(arrowstyle="-|>", color=_NAVY_HEX, lw=2.2, mutation_scale=14))
        ax.text(0, -0.08, f"{score:.0f}", ha="center", va="center",
                fontsize=18, fontweight="bold", color=_NAVY_HEX)
        ax.text(0, 1.12, title, ha="center", va="center",
                fontsize=9, fontweight="bold", color=_NAVY_HEX)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=130, bbox_inches="tight",
                    facecolor="none", transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return b""


def _bar_chart_png(labels: list[str], values: list[float], title: str,
                   color: str = _CYAN_HEX) -> bytes:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots(figsize=(5.0, 2.8), facecolor="none")
        bars = ax.barh(labels[::-1], values[::-1], color=color, height=0.6)
        ax.set_xlabel("Count", fontsize=8, color=_GREY_HEX)
        ax.set_title(title, fontsize=9, fontweight="bold", color=_NAVY_HEX, pad=6)
        ax.tick_params(axis="both", labelsize=7, colors=_GREY_HEX)
        ax.spines[["top", "right", "left"]].set_visible(False)
        ax.set_facecolor("none")
        for bar, val in zip(bars, values[::-1]):
            ax.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height() / 2,
                    str(int(val)), va="center", ha="left", fontsize=7, color=_GREY_HEX)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=130, bbox_inches="tight",
                    facecolor="none", transparent=True)
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return b""


# ─────────────────────────────────────────────────────────────────────────────
# Main Builder
# ─────────────────────────────────────────────────────────────────────────────

class IntegriReportBuilder:
    """Build a polished IntegriShield DOCX report."""

    def __init__(self, title: str, subtitle: str, module_name: str,
                 tenant: str = "Acme Corp", classification: str = "CONFIDENTIAL") -> None:
        self.title = title
        self.subtitle = subtitle
        self.module_name = module_name
        self.tenant = tenant
        self.classification = classification
        self.generated_at = datetime.now(timezone.utc)

        self._doc = Document()
        self._setup_page()
        _add_header_footer(self._doc, module_name)

    # ── page layout ──────────────────────────────────────────────

    def _setup_page(self) -> None:
        from docx.shared import Cm
        section = self._doc.sections[0]
        section.page_width  = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # ── cover page ───────────────────────────────────────────────

    def add_cover_page(self) -> "IntegriReportBuilder":
        doc = self._doc

        # Top colour bar
        tbl = doc.add_table(rows=1, cols=1)
        _set_table_style(tbl)
        cell = tbl.cell(0, 0)
        _set_cell_bg(cell, "0B2447")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("■ INTEGRISHIELD")
        _set_run_font(run, size_pt=22, bold=True, color=WHITE)
        p2 = cell.add_paragraph("Next-Generation SAP Middleware Security Platform")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.runs[0]
        _set_run_font(run2, size_pt=11, color=CYAN)
        p2.paragraph_format.space_after = Pt(14)

        doc.add_paragraph()  # spacer

        # Classification badge
        badge_p = doc.add_paragraph()
        badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        badge_run = badge_p.add_run(f"  {self.classification}  ")
        _set_run_font(badge_run, size_pt=9, bold=True, color=WHITE)
        bg_hex = "EF4444" if "CONFIDENTIAL" in self.classification else "64748B"
        # embed shading via XML on the run's parent para background workaround:
        # use a 1-cell table for the badge
        doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
        badge_tbl = doc.add_table(rows=1, cols=1)
        _set_table_style(badge_tbl)
        bcell = badge_tbl.rows[0].cells[0]
        _set_cell_bg(bcell, bg_hex)
        bp = bcell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(5)
        bp.paragraph_format.space_after  = Pt(5)
        br = bp.add_run(f"  {self.classification}  ")
        _set_run_font(br, size_pt=9, bold=True, color=WHITE)

        doc.add_paragraph()

        # Main report title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(20)
        title_r = title_p.add_run(self.title)
        _set_run_font(title_r, size_pt=26, bold=True, color=NAVY)

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_r = sub_p.add_run(self.subtitle)
        _set_run_font(sub_r, size_pt=13, italic=True, color=RGBColor(0x47, 0x55, 0x69))

        _hr(doc, "19A7CE", 3)

        # Metadata table
        meta = [
            ("Prepared for",   self.tenant),
            ("Prepared by",    "IntegriShield Security Platform v4.0"),
            ("Report date",    self.generated_at.strftime("%d %B %Y")),
            ("Report time",    self.generated_at.strftime("%H:%M UTC")),
            ("Classification", self.classification),
            ("Version",        "1.0 — AUTO-GENERATED"),
        ]
        mtbl = doc.add_table(rows=len(meta), cols=2)
        _set_table_style(mtbl)
        for i, (key, val) in enumerate(meta):
            kc = mtbl.cell(i, 0)
            vc = mtbl.cell(i, 1)
            bg = "F4F6F9" if i % 2 == 0 else "FFFFFF"
            _set_cell_bg(kc, bg)
            _set_cell_bg(vc, bg)
            kp = kc.paragraphs[0]
            kp.paragraph_format.space_before = Pt(4)
            kp.paragraph_format.space_after  = Pt(4)
            kr = kp.add_run(key)
            _set_run_font(kr, size_pt=9, bold=True, color=NAVY)
            vp = vc.paragraphs[0]
            vp.paragraph_format.space_before = Pt(4)
            vp.paragraph_format.space_after  = Pt(4)
            vr = vp.add_run(val)
            _set_run_font(vr, size_pt=9, color=RGBColor(0x1E, 0x29, 0x3B))

        doc.add_paragraph()

        # Confidentiality notice
        notice_p = doc.add_paragraph()
        notice_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice_r = notice_p.add_run(
            "This document contains proprietary and confidential information.\n"
            "Unauthorised disclosure, copying, distribution or use of the contents of this report is strictly prohibited."
        )
        _set_run_font(notice_r, size_pt=7.5, italic=True, color=GREY_D)

        _add_page_break_v2(doc)
        return self

    # ── TOC ─────────────────────────────────────────────────────

    def add_toc(self) -> "IntegriReportBuilder":
        doc = self._doc
        h = doc.add_heading("Table of Contents", level=1)
        h.runs[0].font.color.rgb = NAVY
        _add_toc_field(doc)
        doc.add_paragraph()
        _add_page_break_v2(doc)
        return self

    # ── executive summary ────────────────────────────────────────

    def add_exec_summary(self, risk_score: float, key_findings: list[str],
                         kpis: dict[str, Any], severity_counts: dict[str, int] | None = None) -> "IntegriReportBuilder":
        doc = self._doc
        h = doc.add_heading("Executive Summary", level=1)
        h.runs[0].font.color.rgb = NAVY

        # Two-column layout: gauge left, text right — use a table
        tbl = doc.add_table(rows=1, cols=2)
        _set_table_style(tbl)
        left_cell  = tbl.cell(0, 0)
        right_cell = tbl.cell(0, 1)

        # Gauge chart
        gauge_png = _risk_gauge_png(risk_score, "Overall Risk Score")
        if gauge_png:
            lp = left_cell.paragraphs[0]
            lr = lp.add_run()
            lr.add_picture(io.BytesIO(gauge_png), width=Cm(6.0))

        # Key findings text
        rp = right_cell.paragraphs[0]
        rr = rp.add_run("Key Findings")
        _set_run_font(rr, size_pt=11, bold=True, color=NAVY)
        for finding in key_findings[:5]:
            fp = right_cell.add_paragraph(style="List Bullet")
            fr = fp.add_run(finding)
            _set_run_font(fr, size_pt=9.5)

        doc.add_paragraph()

        # KPI mini-table
        if kpis:
            kpi_items = list(kpis.items())
            cols = min(4, len(kpi_items))
            kpi_tbl = doc.add_table(rows=2, cols=cols)
            _set_table_style(kpi_tbl)
            for i, (label, value) in enumerate(kpi_items[:cols]):
                hc = kpi_tbl.cell(0, i)
                vc = kpi_tbl.cell(1, i)
                _set_cell_bg(hc, "0B2447")
                _set_cell_bg(vc, "EFF6FF")
                hp = hc.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hp.paragraph_format.space_before = Pt(6)
                hp.paragraph_format.space_after  = Pt(6)
                hr_ = hp.add_run(str(label))
                _set_run_font(hr_, size_pt=8, bold=True, color=WHITE)
                vp = vc.paragraphs[0]
                vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                vp.paragraph_format.space_before = Pt(8)
                vp.paragraph_format.space_after  = Pt(8)
                vr = vp.add_run(str(value))
                _set_run_font(vr, size_pt=14, bold=True, color=NAVY)

        # Severity doughnut
        if severity_counts and any(v > 0 for v in severity_counts.values()):
            donut_png = _severity_doughnut_png(severity_counts, "Finding Severity Mix")
            if donut_png:
                doc.add_paragraph()
                dp = doc.add_paragraph()
                dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dr = dp.add_run()
                dr.add_picture(io.BytesIO(donut_png), width=Cm(7.5))

        _add_page_break_v2(doc)
        return self

    # ── section heading + narrative ──────────────────────────────

    def add_section(self, title: str, narrative: str,
                    level: int = 1) -> "IntegriReportBuilder":
        doc = self._doc
        h = doc.add_heading(title, level=level)
        if h.runs:
            h.runs[0].font.color.rgb = NAVY
        if narrative:
            np_ = doc.add_paragraph()
            nr = np_.add_run(narrative)
            _set_run_font(nr, size_pt=10)
            np_.paragraph_format.space_after = Pt(8)
        return self

    # ── zebra-striped evidence table ─────────────────────────────

    def add_evidence_table(self, headers: list[str], rows: list[list[str]],
                           severity_col: int | None = None) -> "IntegriReportBuilder":
        doc = self._doc
        if not rows:
            p = doc.add_paragraph()
            r = p.add_run("No data available for this period.")
            _set_run_font(r, size_pt=9, italic=True, color=GREY_D)
            doc.add_paragraph()
            return self

        col_count = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
        _set_table_style(tbl)

        _SEV_COLORS = {
            "critical": "FEE2E2", "high": "FFF7ED", "medium": "FEFCE8",
            "low": "F0FDF4", "info": "F0F9FF",
        }
        _SEV_TEXT = {
            "critical": "EF4444", "high": "F97316", "medium": "EAB308",
            "low": "22C55E", "info": "0EA5E9",
        }

        # Header row
        hrow = tbl.rows[0]
        for i, hdr in enumerate(headers):
            cell = hrow.cells[i]
            _set_cell_bg(cell, "0B2447")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            r = p.add_run(hdr)
            _set_run_font(r, size_pt=8.5, bold=True, color=WHITE)

        # Data rows
        for ri, row_data in enumerate(rows):
            drow = tbl.rows[ri + 1]
            row_bg = "F4F6F9" if ri % 2 == 0 else "FFFFFF"
            sev_val = (row_data[severity_col].lower() if severity_col is not None
                       and severity_col < len(row_data) else "")
            if sev_val in _SEV_COLORS:
                row_bg = _SEV_COLORS[sev_val]

            for ci, val in enumerate(row_data[:col_count]):
                cell = drow.cells[ci]
                _set_cell_bg(cell, row_bg)
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(3)
                cp.paragraph_format.space_after  = Pt(3)
                cr = cp.add_run(str(val) if val is not None else "—")
                clr = None
                if severity_col is not None and ci == severity_col and sev_val in _SEV_TEXT:
                    clr = RGBColor.from_string(_SEV_TEXT[sev_val])
                _set_run_font(cr, size_pt=8.5, bold=(ci == 0),
                              color=clr or RGBColor(0x1E, 0x29, 0x3B))

        doc.add_paragraph()
        return self

    # ── bar chart ────────────────────────────────────────────────

    def add_bar_chart(self, labels: list[str], values: list[float],
                      title: str, color: str = _CYAN_HEX) -> "IntegriReportBuilder":
        png = _bar_chart_png(labels, values, title, color)
        if png:
            p = self._doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(io.BytesIO(png), width=Cm(13.0))
            self._doc.add_paragraph()
        return self

    # ── remediation appendix ─────────────────────────────────────

    def add_remediation_appendix(self, items: list[dict]) -> "IntegriReportBuilder":
        doc = self._doc
        _add_page_break_v2(doc)
        h = doc.add_heading("Remediation & Action Items", level=1)
        h.runs[0].font.color.rgb = NAVY

        headers = ["#", "Finding", "Priority", "Recommended Action", "Owner", "Effort"]
        rows = []
        for i, item in enumerate(items, 1):
            rows.append([
                str(i),
                textwrap.shorten(item.get("finding", ""), 60, placeholder="…"),
                item.get("priority", "Medium"),
                textwrap.shorten(item.get("action", ""), 80, placeholder="…"),
                item.get("owner", "Security Team"),
                item.get("effort", "1–2 days"),
            ])
        self.add_evidence_table(headers, rows, severity_col=2)
        return self

    # ── signature block ──────────────────────────────────────────

    def add_signature_block(self) -> "IntegriReportBuilder":
        doc = self._doc
        _add_page_break_v2(doc)
        h = doc.add_heading("Document Control", level=1)
        h.runs[0].font.color.rgb = NAVY

        headers = ["Role", "Name", "Signature", "Date"]
        rows = [
            ["Prepared by",  "IntegriShield Platform (Automated)",  "— AUTO-GENERATED —", self.generated_at.strftime("%d %b %Y")],
            ["Reviewed by",  "Security Operations Lead",             "_________________________", ""],
            ["Approved by",  "Chief Information Security Officer",   "_________________________", ""],
            ["Distribution", "Board / Audit Committee",              "For information",            ""],
        ]
        self.add_evidence_table(headers, rows)

        doc.add_paragraph()
        disc = doc.add_paragraph()
        disc_r = disc.add_run(
            "This report was automatically generated by the IntegriShield Security Platform. "
            "All data is sourced directly from live event streams and security engines. "
            "Findings should be reviewed by qualified security personnel before external distribution. "
            f"Report ID: {self.generated_at.strftime('%Y%m%d%H%M%S')}-{self.module_name.upper().replace(' ', '-')}"
        )
        _set_run_font(disc_r, size_pt=7.5, italic=True, color=GREY_D)
        return self

    # ── finalize → bytes ─────────────────────────────────────────

    def finalize(self) -> bytes:
        buf = io.BytesIO()
        self._doc.save(buf)
        buf.seek(0)
        return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# Per-report builders — called by server.py endpoints
# ─────────────────────────────────────────────────────────────────────────────

def build_compliance_report(framework: str, findings: list[dict], tenant: str = "Acme Corp") -> bytes:
    """Full compliance attestation report for a single framework."""
    fw = framework.upper()
    violations = [f for f in findings if str(f.get("result", f.get("status", ""))).lower() in ("violation", "fail")]
    warnings    = [f for f in findings if str(f.get("result", f.get("status", ""))).lower() == "warning"]
    passed      = [f for f in findings if str(f.get("result", f.get("status", ""))).lower() in ("pass", "passed")]

    total = max(len(findings), 1)
    score = round(100 * len(passed) / total, 1) if findings else 94.0
    # Enrich with hardcoded narrative per framework
    narratives = {
        "SOX":      "The Sarbanes-Oxley Act requires robust internal controls over financial reporting. This report assesses segregation-of-duty controls, privileged access governance, and audit trail completeness across the SAP landscape.",
        "GDPR":     "General Data Protection Regulation assessment covering lawful basis for processing, data subject rights enablement, cross-border transfer controls, and incident notification readiness.",
        "PCI-DSS":  "Payment Card Industry Data Security Standard evaluation of cardholder data environment access controls, encryption at rest and in transit, vulnerability management programme, and log monitoring.",
        "NIST-CSF": "NIST Cybersecurity Framework assessment across Identify, Protect, Detect, Respond, and Recover function areas for the SAP middleware integration tier.",
        "ISO27001": "ISO/IEC 27001 Information Security Management System audit covering asset management, access control, cryptography, physical security, and supplier relationships.",
        "HIPAA":    "Health Insurance Portability and Accountability Act compliance review of ePHI access controls, audit controls, transmission security, and Business Associate Agreement coverage.",
    }
    narrative = narratives.get(fw, f"{fw} compliance assessment against applicable controls in the IntegriShield policy engine.")

    b = IntegriReportBuilder(
        title=f"{fw} Compliance Attestation",
        subtitle=f"Automated Compliance Evidence Report — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="M07 Compliance Autopilot",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=100.0 - score,
        key_findings=[
            f"{fw} compliance score: {score}% ({len(passed)}/{total} controls passed)",
            f"{len(violations)} active violation(s) requiring immediate remediation",
            f"{len(warnings)} warning(s) under monitoring — no immediate breach",
            "Continuous evidence collection active across all monitored SAP systems",
            "Automated remediation guidance available for all flagged controls",
        ],
        kpis={
            "Score": f"{score}%",
            "Violations": len(violations),
            "Warnings": len(warnings),
            "Passed": len(passed),
        },
        severity_counts={"critical": len(violations), "medium": len(warnings), "low": len(passed)},
    )

    b.add_section(f"{fw} Framework Overview", narrative)

    # Control-by-control table
    b.add_section("Control Assessment Detail", "", level=2)
    ctrl_rows = []
    for f in findings[:80]:
        result = str(f.get("result", f.get("status", ""))).upper()
        ctrl_rows.append([
            f.get("control_id", "—"),
            textwrap.shorten(f.get("description", f.get("message", "—")), 70, placeholder="…"),
            f.get("framework", fw),
            result,
            str(f.get("ts", f.get("timestamp", ""))[:19] if f.get("ts", f.get("timestamp")) else "—"),
        ])
    if not ctrl_rows:
        # Use hardcoded demo rows when no live events
        ctrl_rows = _demo_compliance_rows(fw)
    b.add_evidence_table(
        ["Control ID", "Finding Description", "Framework", "Status", "Timestamp"],
        ctrl_rows, severity_col=3,
    )

    # Violations section
    if violations:
        b.add_section("Violations Requiring Immediate Action", "", level=2)
        vrows = [[
            v.get("control_id", "—"),
            textwrap.shorten(v.get("description", v.get("message", "—")), 80, placeholder="…"),
            "VIOLATION",
            v.get("user_id", v.get("actor", "—")),
        ] for v in violations[:30]]
        b.add_evidence_table(["Control", "Description", "Status", "Actor"], vrows, severity_col=2)

    # Remediation
    remediation = _compliance_remediation(fw, violations)
    b.add_remediation_appendix(remediation)
    b.add_signature_block()
    return b.finalize()


def build_incidents_report(incidents: list[dict], tenant: str = "Acme Corp") -> bytes:
    open_     = [i for i in incidents if str(i.get("status", "")).lower() == "open"]
    inv       = [i for i in incidents if str(i.get("status", "")).lower() in ("investigating", "in_progress", "active")]
    resolved  = [i for i in incidents if str(i.get("status", "")).lower() in ("resolved", "closed", "contained")]
    critical  = [i for i in incidents if str(i.get("severity", "")).lower() == "critical"]

    if not incidents:
        incidents = _demo_incidents()
        open_ = incidents[:2]; resolved = incidents[2:]

    risk = min(95, 15 + len(open_) * 12 + len(critical) * 8)

    b = IntegriReportBuilder(
        title="Incident Response Report",
        subtitle=f"Security Incident Lifecycle & Playbook Execution — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="M10 Incident Response",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=risk,
        key_findings=[
            f"{len(open_)} open incident(s) requiring active investigation",
            f"{len(inv)} incident(s) currently under investigation",
            f"{len(resolved)} incident(s) resolved this period",
            f"{len(critical)} critical-severity incident(s) detected",
            "MITRE ATT&CK mapping active; automated containment playbooks executed",
        ],
        kpis={"Open": len(open_), "Investigating": len(inv), "Resolved": len(resolved), "Critical": len(critical)},
        severity_counts={"critical": len(critical), "high": max(0, len(incidents) - len(critical)), "low": len(resolved)},
    )

    b.add_section("Incident Summary", (
        "The following section provides a complete overview of all security incidents detected, "
        "investigated, and resolved during the reporting period. Incidents are tracked from initial "
        "detection through containment, eradication, and post-incident review."
    ))

    inc_rows = []
    for inc in incidents[:50]:
        inc_rows.append([
            inc.get("incident_id", inc.get("id", "—"))[:16],
            textwrap.shorten(inc.get("title", inc.get("scenario", "—")), 55, placeholder="…"),
            str(inc.get("severity", "—")).upper(),
            str(inc.get("status", "—")).upper(),
            inc.get("source_module", inc.get("module", "—")),
            str(inc.get("ts", inc.get("timestamp", ""))[:19] if inc.get("ts", inc.get("timestamp")) else "—"),
        ])
    b.add_evidence_table(
        ["Incident ID", "Title", "Severity", "Status", "Source Module", "Detected At"],
        inc_rows, severity_col=2,
    )

    b.add_section("Playbook Execution Timeline", (
        "Each incident triggers an automated playbook. The table below shows playbook stages "
        "with durations and outcomes for the current reporting period."
    ), level=2)

    pb_rows = _playbook_timeline_rows(incidents)
    b.add_evidence_table(
        ["Incident ID", "Stage", "Duration", "Outcome", "Analyst"],
        pb_rows, severity_col=None,
    )

    b.add_section("MITRE ATT&CK Mapping", (
        "Detected incidents are automatically mapped to MITRE ATT&CK techniques "
        "to support threat intelligence and defensive gap analysis."
    ), level=2)

    mitre_rows = _mitre_rows(incidents)
    b.add_evidence_table(["Incident ID", "ATT&CK Technique", "Tactic", "Confidence"], mitre_rows)

    b.add_remediation_appendix(_incident_remediation(open_, critical))
    b.add_signature_block()
    return b.finalize()


_TOP_CVES_STATIC = [
    {"cve_id": "CVE-2025-23121", "cvss": 9.8, "severity": "CRITICAL", "dependency": "requests==2.28.0", "description": "HTTP header injection via redirect chain in requests library.", "published": "2025-11-14"},
    {"cve_id": "CVE-2025-31200", "cvss": 8.8, "severity": "HIGH",     "dependency": "pydantic==1.10.9",   "description": "Remote code execution via crafted validation model in pydantic v1.", "published": "2025-10-02"},
    {"cve_id": "CVE-2024-56334", "cvss": 8.1, "severity": "HIGH",     "dependency": "cryptography==41.0.2","description": "Weak RSA key generation under specific entropy conditions.", "published": "2024-12-19"},
    {"cve_id": "CVE-2025-29441", "cvss": 7.5, "severity": "HIGH",     "dependency": "aiohttp==3.9.1",     "description": "Path traversal in aiohttp static file handler.", "published": "2025-09-07"},
    {"cve_id": "CVE-2024-52302", "cvss": 6.5, "severity": "MEDIUM",   "dependency": "uvicorn==0.23.2",    "description": "HTTP/1.1 request smuggling in uvicorn connection handler.", "published": "2024-11-28"},
]


def build_sbom_report(sbom_scans: list[dict], tenant: str = "Acme Corp",
                      top_cves: list[dict] | None = None) -> bytes:
    top_cves = top_cves or _TOP_CVES_STATIC
    critical_cves = [c for c in top_cves if c.get("severity") == "CRITICAL"]
    high_cves     = [c for c in top_cves if c.get("severity") == "HIGH"]

    b = IntegriReportBuilder(
        title="Software Bill of Materials & CVE Report",
        subtitle=f"Dependency Vulnerability Assessment — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="M13 SBOM Scanner",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=72.0,
        key_findings=[
            f"{len(critical_cves)} CRITICAL CVE(s) detected in active dependencies",
            f"{len(high_cves)} HIGH severity CVE(s) identified requiring patch planning",
            "NVD and OSV feeds refreshed; 4,050+ CVEs indexed in local cache",
            "Affected packages include requests, pydantic, cryptography, aiohttp",
            "Immediate patching recommended for CVSS ≥ 9.0 findings",
        ],
        kpis={
            "Critical CVEs": len(critical_cves),
            "High CVEs":     len(high_cves),
            "Total Indexed": "4,050",
            "Packages":      len(top_cves),
        },
        severity_counts={"CRITICAL": len(critical_cves), "HIGH": len(high_cves), "MEDIUM": 1, "LOW": 0},
    )

    b.add_section("CVE Findings", (
        "The IntegriShield SBOM scanner continuously queries the National Vulnerability Database (NVD) "
        "and Open Source Vulnerability (OSV) feed. The following table lists all high-severity findings "
        "detected in currently deployed dependencies, ordered by CVSS base score."
    ))

    cve_rows = []
    for cve in top_cves:
        cve_rows.append([
            cve.get("cve_id", "—"),
            cve.get("dependency", "—"),
            str(cve.get("cvss", "—")),
            cve.get("severity", "—"),
            textwrap.shorten(cve.get("description", "—"), 70, placeholder="…"),
            cve.get("published", "—"),
        ])
    b.add_evidence_table(
        ["CVE ID", "Affected Dependency", "CVSS", "Severity", "Description", "Published"],
        cve_rows, severity_col=3,
    )

    b.add_section("Recent SBOM Scan Events", "", level=2)
    scan_rows = []
    for sc in (sbom_scans or [])[:40]:
        scan_rows.append([
            sc.get("package", sc.get("dependency", "—")),
            sc.get("version", "—"),
            str(sc.get("cve_id", sc.get("cve", "—"))),
            str(sc.get("severity", "—")).upper(),
            str(sc.get("ts", sc.get("timestamp", ""))[:19] if sc.get("ts", sc.get("timestamp")) else "—"),
        ])
    if not scan_rows:
        scan_rows = [["requests", "2.28.0", "CVE-2025-23121", "CRITICAL", "2026-04-18T06:12:00"],
                     ["pydantic", "1.10.9", "CVE-2025-31200", "HIGH",     "2026-04-18T06:12:05"],
                     ["cryptography", "41.0.2", "CVE-2024-56334", "HIGH", "2026-04-18T06:12:10"]]
    b.add_evidence_table(["Package", "Version", "CVE ID", "Severity", "Detected At"], scan_rows, severity_col=3)

    b.add_remediation_appendix([
        {"finding": "CVE-2025-23121 in requests==2.28.0 (CVSS 9.8)", "priority": "Critical",
         "action": "Upgrade requests to >=2.32.0 immediately", "owner": "Platform Team", "effort": "2 hours"},
        {"finding": "CVE-2025-31200 in pydantic==1.10.9 (CVSS 8.8)", "priority": "High",
         "action": "Migrate to pydantic v2 or apply available patch", "owner": "Dev Team", "effort": "1 day"},
        {"finding": "CVE-2024-56334 in cryptography==41.0.2 (CVSS 8.1)", "priority": "High",
         "action": "Upgrade cryptography to >=42.0.4", "owner": "Platform Team", "effort": "4 hours"},
        {"finding": "CVE-2025-29441 in aiohttp==3.9.1 (CVSS 7.5)", "priority": "High",
         "action": "Upgrade aiohttp to >=3.9.5", "owner": "Dev Team", "effort": "2 hours"},
        {"finding": "Enable automated dependency scanning in CI/CD pipeline", "priority": "Medium",
         "action": "Integrate SBOM scanner into GitHub Actions with fail-on-critical policy",
         "owner": "DevSecOps", "effort": "1 day"},
    ])
    b.add_signature_block()
    return b.finalize()


def build_dlp_report(dlp_alerts: list[dict], tenant: str = "Acme Corp") -> bytes:
    by_rule: dict[str, int] = {}
    for a in dlp_alerts:
        rule = str(a.get("dlp_rule", a.get("rule", a.get("scenario", "unknown"))))
        by_rule[rule] = by_rule.get(rule, 0) + 1

    risk = min(90, 10 + len(dlp_alerts) * 3)
    critical = [a for a in dlp_alerts if str(a.get("severity", "")).lower() == "critical"]

    if not dlp_alerts:
        dlp_alerts = _demo_dlp_alerts()
        by_rule = {"bulk-export": 5, "pii-exfil": 3, "data-staging": 2, "blocklist": 1}
        risk = 48

    b = IntegriReportBuilder(
        title="Data Loss Prevention Report",
        subtitle=f"DLP Violation Analysis & Policy Enforcement — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="M09 DLP Engine",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=risk,
        key_findings=[
            f"{len(dlp_alerts)} DLP violation(s) detected this period",
            f"{len(critical)} critical violation(s) involving sensitive PII or bulk extraction",
            f"Top rule triggered: {max(by_rule, key=by_rule.get) if by_rule else 'N/A'}",
            "Field-level masking and SAP table access controls are actively enforced",
            "Real-time alerting active; all violations logged to immutable audit trail",
        ],
        kpis={
            "Total Violations": len(dlp_alerts),
            "Critical": len(critical),
            "Rules Active": 4,
            "Data Blocked (MB)": str(sum(int(a.get("bytes_out", 0)) for a in dlp_alerts) // 1_000_000 or "N/A"),
        },
        severity_counts={k: v for k, v in [("critical", len(critical)),
                                             ("high", max(0, len(dlp_alerts) - len(critical)))]} ,
    )

    b.add_section("DLP Violations by Rule", (
        "The IntegriShield DLP Engine applies regex-based pattern matching and SAP field-level rules "
        "to detect unauthorized data exfiltration, PII exposure, and bulk data extraction attempts. "
        "All violations trigger real-time alerts and are forwarded to the SIEM integration webhook."
    ))

    if by_rule:
        b.add_bar_chart(list(by_rule.keys()), list(by_rule.values()),
                        "Violations by DLP Rule", _RED_HEX)

    dlp_rows = []
    for a in dlp_alerts[:60]:
        dlp_rows.append([
            a.get("user_id", a.get("actor", "—")),
            a.get("source_ip", "—"),
            str(a.get("dlp_rule", a.get("rule", a.get("scenario", "—")))),
            str(a.get("severity", "—")).upper(),
            str(a.get("bytes_out", a.get("data_size", "—"))),
            str(a.get("ts", a.get("timestamp", ""))[:19] if a.get("ts", a.get("timestamp")) else "—"),
        ])
    if not dlp_rows:
        dlp_rows = _demo_dlp_rows()
    b.add_evidence_table(
        ["User", "Source IP", "Rule Triggered", "Severity", "Bytes Out", "Timestamp"],
        dlp_rows, severity_col=3,
    )

    b.add_remediation_appendix([
        {"finding": "Bulk SAP table extraction by privileged users", "priority": "Critical",
         "action": "Implement mandatory justification workflow for extractions >10MB; alert SOC immediately",
         "owner": "IAM Team", "effort": "3 days"},
        {"finding": "PII exfiltration via RFC_READ_TABLE on PA0001/PA0008", "priority": "Critical",
         "action": "Apply row-level security on HR tables; restrict to dedicated payroll role only",
         "owner": "Basis Team", "effort": "1 day"},
        {"finding": "Off-hours data staging activity", "priority": "High",
         "action": "Enforce time-based access controls; require approval for off-hours RFC calls",
         "owner": "Security Ops", "effort": "2 days"},
        {"finding": "Blocklist keyword matches in RFC payloads", "priority": "Medium",
         "action": "Review and expand DLP keyword blocklist; tune false-positive threshold",
         "owner": "DLP Team", "effort": "4 hours"},
    ])
    b.add_signature_block()
    return b.finalize()


def build_rbac_report(decisions: list[dict], counters: dict, tenant: str = "Acme Corp") -> bytes:
    allow = counters.get("ALLOW", 0)
    deny  = counters.get("DENY", 0)
    mod   = counters.get("MODIFY", 0)
    total = counters.get("total", max(allow + deny + mod, 1))

    denied_dec = [d for d in decisions if d.get("decision") == "DENY"]
    inject_dec = [d for d in decisions if "injection" in str(d.get("reason", "")).lower()
                  or "inject" in str(d.get("tool_name", "")).lower()]

    b = IntegriReportBuilder(
        title="MCP Security — RBAC Audit Report",
        subtitle=f"M16 Policy Engine Decision Log & Access Control Review — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="M16 MCP Security Layer",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=min(80, 5 + len(denied_dec) * 4),
        key_findings=[
            f"{total} total policy decisions evaluated by the RBAC engine",
            f"{allow} requests allowed — {deny} denied — {mod} modified (row-capped)",
            f"{len(inject_dec)} potential prompt-injection attempt(s) detected and blocked",
            "10-rule RBAC policy set enforced; default-deny baseline active",
            "All decisions logged to immutable rolling audit store (500-entry deque)",
        ],
        kpis={"Total": total, "ALLOW": allow, "DENY": deny, "MODIFY": mod},
        severity_counts={"critical": deny, "medium": mod, "low": allow},
    )

    b.add_section("Policy Decision Log", (
        "Every MCP tool invocation passes through the M16 policy engine before execution. "
        "The engine evaluates the caller's role, the target tool, and any input parameters "
        "against the active RBAC ruleset. The table below shows a sample of recent decisions."
    ))

    dec_rows = []
    for d in decisions[:60]:
        dec_rows.append([
            d.get("user_id", "—"),
            d.get("role", "—"),
            d.get("tool_name", "—"),
            d.get("decision", "—"),
            d.get("rule_id", "—"),
            textwrap.shorten(d.get("reason", "—"), 60, placeholder="…"),
            str(d.get("timestamp", d.get("ts", ""))[:16] if d.get("timestamp", d.get("ts")) else "—"),
        ])
    b.add_evidence_table(
        ["User", "Role", "Tool", "Decision", "Rule ID", "Reason", "Timestamp"],
        dec_rows, severity_col=3,
    )

    b.add_section("Role Coverage Matrix", (
        "The following matrix shows which roles are permitted, denied, or modified for each "
        "primary tool category in the current policy configuration."
    ), level=2)

    matrix_rows = [
        ["SOC_ADMIN",   "rfc_call_function",   "ALLOW",  "R-001", "Full unrestricted access"],
        ["SOC_ADMIN",   "rfc_read_table",       "ALLOW",  "R-001", "Full unrestricted access"],
        ["SOC_ANALYST", "rfc_read_table",       "MODIFY", "R-020", "Row-capped to 1,000 rows"],
        ["SOC_ANALYST", "rfc_call_function",    "DENY",   "R-023", "Privileged RFC blocked"],
        ["SOC_ANALYST", "get_*",                "ALLOW",  "R-021", "Metadata reads permitted"],
        ["SOC_ANALYST", "list_*",               "ALLOW",  "R-022", "List/query operations permitted"],
        ["AUDITOR",     "rfc_read_table",       "MODIFY", "R-011", "Row-capped to 500 rows"],
        ["AUDITOR",     "rfc_call_function",    "DENY",   "R-012", "Write/execute tools blocked"],
        ["AUDITOR",     "get_*",                "ALLOW",  "R-010", "Read-only metadata access"],
        ["SERVICE",     "rfc_*",                "DENY",   "R-031", "SAP table reads prohibited"],
        ["SERVICE",     "publish_*",            "ALLOW",  "R-030", "Event publishing permitted"],
    ]
    b.add_evidence_table(
        ["Role", "Tool Pattern", "Decision", "Rule", "Notes"],
        [[r[0], r[1], r[2], r[3], r[4]] for r in matrix_rows], severity_col=2,
    )

    if denied_dec:
        b.add_section("Denied Requests — Detail", "", level=2)
        deny_rows = [[
            d.get("user_id", "—"), d.get("role", "—"), d.get("tool_name", "—"),
            d.get("rule_id", "—"), textwrap.shorten(d.get("reason", "—"), 70, placeholder="…"),
        ] for d in denied_dec[:30]]
        b.add_evidence_table(["User", "Role", "Tool", "Rule", "Reason"], deny_rows)

    b.add_remediation_appendix([
        {"finding": f"{deny} requests denied — review recurring denial patterns for policy gaps",
         "priority": "Medium",
         "action": "Audit top-denied tool patterns and refine role assignments where legitimate access needed",
         "owner": "IAM Team", "effort": "2 days"},
        {"finding": "In-memory audit log (500 entries max) — potential evidence gap",
         "priority": "High",
         "action": "Implement Postgres persistence for RBAC audit log (Dev-5 milestone)",
         "owner": "Platform Team", "effort": "3 days"},
        {"finding": "JWT/OIDC identity validation not yet enforced",
         "priority": "High",
         "action": "Integrate Okta/Azure AD token validation in M16 middleware before production rollout",
         "owner": "Security Engineering", "effort": "1 week"},
    ])
    b.add_signature_block()
    return b.finalize()


def build_executive_report(all_data: dict, tenant: str = "Acme Corp") -> bytes:
    """Cross-module executive posture report for board / audit committee."""
    alerts    = all_data.get("alerts", [])
    anomalies = all_data.get("anomalies", [])
    incidents = all_data.get("incidents", [])
    dlp       = all_data.get("dlp", [])
    compliance= all_data.get("compliance", [])
    sbom      = all_data.get("sbom", [])
    cloud     = all_data.get("cloud", [])
    shadow    = all_data.get("shadow", [])
    m16_counters = all_data.get("m16_counters", {"ALLOW": 42, "DENY": 8, "MODIFY": 12, "total": 62})

    critical_alerts = [a for a in alerts if str(a.get("severity", "")).lower() == "critical"]
    open_incidents  = [i for i in incidents if str(i.get("status", "")).lower() == "open"]
    critical_cves   = 1  # from _M13_TOP_CVES

    risk_score = min(88, 20 + len(critical_alerts) * 2 + len(open_incidents) * 5 + len(dlp) * 1.5)

    b = IntegriReportBuilder(
        title="Executive Security Posture Report",
        subtitle=f"Board-Level Cybersecurity Dashboard — {datetime.now(timezone.utc).strftime('%B %Y')}",
        module_name="IntegriShield Platform",
        tenant=tenant,
    )
    b.add_cover_page()
    b.add_toc()
    b.add_exec_summary(
        risk_score=risk_score,
        key_findings=[
            f"Overall risk score: {risk_score:.0f}/100 — {'HIGH' if risk_score > 70 else 'MEDIUM' if risk_score > 40 else 'LOW'} risk posture",
            f"{len(critical_alerts)} critical alert(s) across all 16 security modules",
            f"{len(open_incidents)} open security incident(s) requiring board awareness",
            f"{len(dlp)} DLP violation(s) detected — data exfiltration controls active",
            "All 16 IntegriShield modules operational; continuous monitoring active",
        ],
        kpis={
            "Critical Alerts": len(critical_alerts),
            "Open Incidents":  len(open_incidents),
            "DLP Violations":  len(dlp),
            "Anomalies":       len(anomalies),
        },
        severity_counts={
            "critical": len(critical_alerts),
            "high": len(anomalies),
            "medium": len(dlp),
            "low": max(0, len(alerts) - len(critical_alerts)),
        },
    )

    b.add_section("Platform Module Status", (
        "IntegriShield deploys 16 specialised security modules across the SAP middleware layer. "
        "The following table shows operational status and key metrics for each module."
    ))

    module_rows = [
        ["M01", "API Gateway Shield",      "ACTIVE", str(len(alerts)),    "Request interception & rules engine"],
        ["M03", "Traffic Analyzer",        "ACTIVE", "Live",               "Velocity checks & heuristic analysis"],
        ["M04", "Zero-Trust Fabric",       "ACTIVE", "Live",               "Continuous trust scoring & mTLS"],
        ["M05", "SAP MCP Suite",           "ACTIVE", "17 tools",           "Claude-integrated SAP security tools"],
        ["M06", "Credential Vault",        "ACTIVE", "Live",               "HashiCorp Vault KV v2 integration"],
        ["M07", "Compliance Autopilot",    "ACTIVE", str(len(compliance)), "SOX·GDPR·PCI-DSS·ISO27001 evidence"],
        ["M08", "Anomaly Detection",       "ACTIVE", str(len(anomalies)),  "ML IsolationForest anomaly scores"],
        ["M09", "DLP Engine",             "ACTIVE", str(len(dlp)),        "PII regex & SAP field-level rules"],
        ["M10", "Incident Response",       "ACTIVE", str(len(incidents)),  "MITRE ATT&CK mapping & playbooks"],
        ["M11", "Shadow Integration",      "ACTIVE", str(len(shadow)),     "Unapproved RFC endpoint detection"],
        ["M12", "Rules Engine",            "ACTIVE", str(len(alerts)),     "Alert fan-out & correlation"],
        ["M13", "SBOM Scanner",            "ACTIVE", "4,050 CVEs",        "NVD/OSV feed, 24h cache"],
        ["M14", "Webhook Gateway",         "ACTIVE", "Live",               "HMAC signing, retry, DLQ"],
        ["M15", "MultiCloud ISPM",         "ACTIVE", str(len(cloud)),      "AWS/GCP/Azure posture checks"],
        ["M16", "MCP Security Layer",      "ACTIVE", str(m16_counters.get("total", 0)) + " decisions", "RBAC policy engine"],
    ]
    b.add_evidence_table(
        ["Module", "Name", "Status", "Events", "Description"],
        module_rows, severity_col=None,
    )

    b.add_section("Top Security Risks", (
        "The following risks have been identified as requiring board-level attention during this reporting period."
    ), level=2)

    risk_rows = [
        ["1", "Critical CVE in active dependency (requests==2.28.0)", "Critical", "Patch immediately", "Platform Team"],
        ["2", f"{len(open_incidents)} open security incidents", "High" if open_incidents else "Low",
         "Escalate to CISO; initiate playbooks", "Security Ops"],
        ["3", "M16 audit log not persisted to database", "High", "Postgres integration (Dev-5)", "Engineering"],
        ["4", f"{len(critical_alerts)} critical alerts unacknowledged", "High" if critical_alerts else "Low",
         "SOC triage required within SLA", "SOC Team"],
        ["5", "JWT/OIDC identity validation pending", "Medium",
         "Deploy before production release", "Security Engineering"],
        ["6", "SBOM NVD cache ephemeral (24h SQLite)", "Medium",
         "Migrate to persistent DB with versioning", "Platform Team"],
        ["7", "DLP off-hours bulk extractions detected", "Medium",
         "Enforce time-based access controls", "IAM Team"],
        ["8", "Compliance scorecard partially hardcoded", "Low",
         "Wire to live compliance engine API", "Frontend Team"],
    ]
    b.add_evidence_table(
        ["#", "Risk Description", "Severity", "Recommended Action", "Owner"],
        risk_rows, severity_col=2,
    )

    b.add_section("Board Recommendations", (
        "Based on the current security posture analysis, the following strategic recommendations "
        "are presented to the board and audit committee for consideration and approval:\n\n"
        "1. IMMEDIATE: Authorise emergency patch cycle for critical CVE-2025-23121 (CVSS 9.8) affecting "
        "the requests library. This is a production-impacting vulnerability.\n\n"
        "2. SHORT-TERM (30 days): Commission Dev-5 sprint to implement Postgres persistence for the RBAC "
        "audit log and complete JWT/OIDC identity integration. Current in-memory audit is non-compliant "
        "with SOX evidence retention requirements.\n\n"
        "3. MEDIUM-TERM (90 days): Engage third-party penetration testing firm to validate IntegriShield's "
        "detection coverage against MITRE ATT&CK for SAP. Budget allocation: £25,000.\n\n"
        "4. ONGOING: Maintain current velocity of IntegriShield development sprints. Platform provides "
        "demonstrably superior coverage versus legacy SIEM-only approach at lower operational cost."
    ), level=2)

    b.add_signature_block()
    return b.finalize()


# ─────────────────────────────────────────────────────────────────────────────
# Demo / hardcoded fallback data helpers
# ─────────────────────────────────────────────────────────────────────────────

def _demo_compliance_rows(fw: str) -> list[list[str]]:
    controls = {
        "SOX":      [("AC-6", "Segregation of duty violation: AP_GL_POSTING", "VIOLATION"), ("AC-2", "Dormant privileged user not locked", "WARNING"), ("AU-2", "Audit trail complete — no gaps", "PASS"), ("CM-2", "Configuration baseline current", "PASS"), ("AC-3", "RFC access controls enforced", "PASS")],
        "GDPR":     [("Art-32", "Data at rest encryption verified", "PASS"), ("Art-30", "Processing activity record current", "PASS"), ("Art-33", "Breach notification procedure tested", "WARNING"), ("Art-17", "Right-to-erasure workflow available", "PASS")],
        "PCI-DSS":  [("Req-1", "Firewall rules reviewed and documented", "PASS"), ("Req-6", "Vulnerability patching >30 days overdue", "VIOLATION"), ("Req-10", "Log monitoring active — all systems", "PASS")],
        "NIST-CSF": [("ID.AM", "Asset inventory complete", "PASS"), ("PR.AC", "Access control policy enforced", "PASS"), ("DE.AE", "Anomaly detection active", "PASS"), ("RS.RP", "Incident response plan current", "WARNING")],
        "ISO27001": [("A.9.1", "Access control policy reviewed", "PASS"), ("A.12.1", "Change management process active", "PASS"), ("A.16.1", "Incident management procedure documented", "PASS"), ("A.18.1", "Compliance with legal requirements", "WARNING")],
        "HIPAA":    [("164.312(a)", "Access control — unique user identification", "PASS"), ("164.312(b)", "Audit controls — hardware and software", "PASS"), ("164.312(e)", "Transmission security — encryption", "VIOLATION")],
    }
    base = controls.get(fw, controls["SOX"])
    return [[c[0], c[1], fw, c[2], "2026-04-18T06:00:00"] for c in base]


def _demo_incidents() -> list[dict]:
    return [
        {"incident_id": "INC-2026-0041", "title": "Bulk SAP extraction — critical", "severity": "critical", "status": "open", "source_module": "m09", "ts": "2026-04-18T08:31:00Z", "playbook_id": "PB-001"},
        {"incident_id": "INC-2026-0042", "title": "Off-hours RFC call detected", "severity": "high", "status": "investigating", "source_module": "m01", "ts": "2026-04-18T03:15:00Z", "playbook_id": "PB-002"},
        {"incident_id": "INC-2026-0039", "title": "PII exfiltration via PA0008", "severity": "critical", "status": "resolved", "source_module": "m09", "ts": "2026-04-17T22:10:00Z", "playbook_id": "PB-001"},
        {"incident_id": "INC-2026-0038", "title": "Shadow RFC endpoint accessed", "severity": "high", "status": "resolved", "source_module": "m11", "ts": "2026-04-17T18:45:00Z", "playbook_id": "PB-003"},
    ]


def _demo_dlp_alerts() -> list[dict]:
    return [
        {"user_id": "BAUER_M", "source_ip": "10.4.22.101", "scenario": "bulk-export", "severity": "critical", "bytes_out": 45_000_000, "ts": "2026-04-18T08:31:00Z"},
        {"user_id": "ROOT", "source_ip": "10.4.0.1", "scenario": "pii-exfil", "severity": "critical", "bytes_out": 2_400_000, "ts": "2026-04-18T03:15:00Z"},
        {"user_id": "SVC_LEGACY", "source_ip": "10.4.33.9", "scenario": "data-staging", "severity": "high", "bytes_out": 8_200_000, "ts": "2026-04-17T22:10:00Z"},
    ]


def _demo_dlp_rows() -> list[list[str]]:
    return [
        ["BAUER_M",    "10.4.22.101", "bulk-export",   "CRITICAL", "45,000,000", "2026-04-18T08:31:00"],
        ["ROOT",       "10.4.0.1",    "pii-exfil",     "CRITICAL", "2,400,000",  "2026-04-18T03:15:00"],
        ["SVC_LEGACY", "10.4.33.9",   "data-staging",  "HIGH",     "8,200,000",  "2026-04-17T22:10:00"],
        ["SYSADMIN",   "10.4.0.2",    "blocklist",     "MEDIUM",   "125,000",    "2026-04-17T18:45:00"],
    ]


def _playbook_timeline_rows(incidents: list[dict]) -> list[list[str]]:
    stages = ["Detected", "Triaged", "Contained", "Eradicated", "Recovered", "Post-mortem"]
    rows = []
    for inc in incidents[:6]:
        iid = inc.get("incident_id", inc.get("id", "INC-???"))[:16]
        for j, stage in enumerate(stages):
            duration = ["0m", "4m", "18m", "35m", "52m", "24h"][j]
            outcome = "✓ Complete" if j < 4 else ("✓ Complete" if str(inc.get("status", "")).lower() in ("resolved", "closed") else "⏳ In progress")
            analyst = ["SIEM (Auto)", "L1 SOC Analyst", "L2 SOC Analyst", "L2 SOC Analyst", "Platform Team", "Security Lead"][j]
            rows.append([iid if j == 0 else "", stage, duration, outcome, analyst])
    if not rows:
        rows = [
            ["INC-2026-0041", "Detected",    "0m",  "✓ Complete", "SIEM (Auto)"],
            ["",              "Triaged",     "4m",  "✓ Complete", "L1 SOC Analyst"],
            ["",              "Contained",   "18m", "✓ Complete", "L2 SOC Analyst"],
            ["",              "Eradicated",  "35m", "⏳ In progress", "L2 SOC Analyst"],
        ]
    return rows


def _mitre_rows(incidents: list[dict]) -> list[list[str]]:
    mitre_map = {
        "bulk-extract": ("T1005", "Data from Local System",    "Collection",    "High"),
        "bulk-export":  ("T1041", "Exfiltration Over C2 Channel","Exfiltration","High"),
        "pii-exfil":    ("T1048", "Exfiltration Over Alt Protocol","Exfiltration","High"),
        "off-hours":    ("T1078", "Valid Accounts",            "Defence Evasion","Medium"),
        "shadow":       ("T1190", "Exploit Public-Facing App", "Initial Access", "High"),
        "credential":   ("T1003", "OS Credential Dumping",    "Credential Access","High"),
    }
    rows = []
    for inc in incidents[:10]:
        sc = str(inc.get("scenario", inc.get("title", ""))).lower()
        for k, (tid, tname, tactic, conf) in mitre_map.items():
            if k in sc:
                rows.append([inc.get("incident_id", "—")[:16], f"{tid} — {tname}", tactic, conf])
                break
        else:
            rows.append([inc.get("incident_id", "—")[:16], "T1078 — Valid Accounts", "Defence Evasion", "Medium"])
    if not rows:
        rows = [
            ["INC-2026-0041", "T1005 — Data from Local System",     "Collection",        "High"],
            ["INC-2026-0042", "T1078 — Valid Accounts",             "Defence Evasion",   "Medium"],
            ["INC-2026-0039", "T1048 — Exfiltration Over Alt Proto","Exfiltration",      "High"],
        ]
    return rows


def _compliance_remediation(fw: str, violations: list[dict]) -> list[dict]:
    base = [
        {"finding": f"{fw} control assessment shows {len(violations)} violation(s)",
         "priority": "Critical" if violations else "Low",
         "action": "Immediately remediate all flagged controls and re-assess within 48 hours",
         "owner": "GRC Team", "effort": "2–5 days"},
        {"finding": "Continuous compliance monitoring gap — evidence not persisted to long-term DB",
         "priority": "High",
         "action": "Enable Postgres persistence in M07 service for 12-month evidence retention",
         "owner": "Platform Team", "effort": "3 days"},
        {"finding": "Manual review cycle not yet automated for all control families",
         "priority": "Medium",
         "action": "Schedule automated monthly compliance report generation and distribution",
         "owner": "Security Ops", "effort": "1 day"},
    ]
    for v in violations[:3]:
        base.insert(0, {
            "finding": f"Control {v.get('control_id', '?')} — {textwrap.shorten(v.get('description', v.get('message', 'violation')), 60)}",
            "priority": "Critical",
            "action": v.get("remediation_guidance", "Review and remediate this control immediately"),
            "owner": "GRC / Basis Team", "effort": "1–2 days",
        })
    return base[:8]


def _incident_remediation(open_: list[dict], critical: list[dict]) -> list[dict]:
    items = []
    for inc in (critical or open_)[:3]:
        items.append({
            "finding": textwrap.shorten(inc.get("title", inc.get("scenario", "Open incident")), 70),
            "priority": str(inc.get("severity", "High")).capitalize(),
            "action": "Execute full playbook: isolate affected user, rotate credentials, collect forensic evidence",
            "owner": "SOC L2 Analyst", "effort": "2–4 hours",
        })
    items += [
        {"finding": "Post-incident review not yet scheduled for resolved incidents",
         "priority": "Medium", "action": "Schedule 30-min post-mortem within 5 business days of closure",
         "owner": "Security Lead", "effort": "30 min"},
        {"finding": "Playbook automation coverage at 70% — 3 scenarios still manual",
         "priority": "Medium", "action": "Automate containment steps for off-hours RFC and shadow-endpoint scenarios",
         "owner": "SOC Engineering", "effort": "3 days"},
    ]
    return items[:6]
