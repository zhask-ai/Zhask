"""Report generator — builds JSON, CSV, and DOCX compliance reports."""

from __future__ import annotations

import csv
import io
import textwrap
import uuid
from collections import OrderedDict
from datetime import datetime, timezone
from typing import Any

from integrishield.m07.models import Framework, ReportRequest


class ReportGenerator:
    """Generates downloadable compliance reports from engine state."""

    def __init__(self, engine) -> None:
        self._engine = engine
        self._reports: OrderedDict[str, dict[str, Any]] = OrderedDict()
        self._max_reports = 50

    def generate(self, req: ReportRequest) -> str:
        """Generate a report and return its report_id."""
        report_id = str(uuid.uuid4())
        now = datetime.now(tz=timezone.utc)

        summary = self._engine.get_summary(req.framework)
        assessments = self._engine.get_assessments(req.framework)

        report_data: dict[str, Any] = {
            "report_id": report_id,
            "framework": req.framework.value,
            "generated_at": now.isoformat(),
            "tenant_id": req.tenant_id,
            "summary": summary.model_dump(),
            "controls": [],
        }

        for assessment in assessments:
            evidence = self._engine.get_evidence(assessment.control_id, limit=50)
            ctrl_def = self._engine._loader.get_control(assessment.control_id)
            report_data["controls"].append(
                {
                    "control_id": assessment.control_id,
                    "title": ctrl_def.title if ctrl_def else assessment.control_id,
                    "status": assessment.status.value,
                    "evidence_count": assessment.evidence_count,
                    "violation_count": assessment.violation_count,
                    "last_assessed_at": assessment.last_assessed_at.isoformat(),
                    "last_violation_at": (
                        assessment.last_violation_at.isoformat()
                        if assessment.last_violation_at
                        else None
                    ),
                    "recent_evidence": [e.model_dump() for e in evidence[:10]],
                    "remediation_guidance": ctrl_def.remediation_guidance if ctrl_def else "",
                }
            )

        self._reports[report_id] = report_data
        if len(self._reports) > self._max_reports:
            self._reports.popitem(last=False)

        return report_id

    def get_json(self, report_id: str) -> dict | None:
        return self._reports.get(report_id)

    def get_csv(self, report_id: str) -> str | None:
        report = self._reports.get(report_id)
        if report is None:
            return None

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow([
            "framework", "control_id", "title", "status",
            "evidence_count", "violation_count",
            "last_assessed_at", "last_violation_at",
        ])
        for ctrl in report.get("controls", []):
            writer.writerow([
                report["framework"],
                ctrl["control_id"],
                ctrl["title"],
                ctrl["status"],
                ctrl["evidence_count"],
                ctrl["violation_count"],
                ctrl["last_assessed_at"],
                ctrl.get("last_violation_at", ""),
            ])
        return output.getvalue()

    def get_docx(self, report_id: str) -> bytes | None:
        """Generate a polished DOCX compliance report using the shared report builder."""
        report = self._reports.get(report_id)
        if report is None:
            return None

        try:
            from integrishield.m07.services._docx_builder import build_m07_docx
            return build_m07_docx(report)
        except Exception:
            return self._get_docx_fallback(report)

    def _get_docx_fallback(self, report: dict) -> bytes:
        """Minimal DOCX fallback when the full builder is unavailable."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document()
            fw  = report.get("framework", "Unknown")
            doc.add_heading(f"{fw} Compliance Report", level=0)
            doc.add_paragraph(f"Generated: {report.get('generated_at', '')}")
            doc.add_paragraph(f"Tenant: {report.get('tenant_id', 'N/A')}")
            summary = report.get("summary", {})
            doc.add_heading("Summary", level=1)
            for k, v in summary.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")
            doc.add_heading("Controls", level=1)
            for ctrl in report.get("controls", []):
                doc.add_paragraph(
                    f"{ctrl['control_id']} — {ctrl['title']} — {ctrl['status'].upper()}",
                    style="List Bullet",
                )
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()
        except Exception:
            return b""
